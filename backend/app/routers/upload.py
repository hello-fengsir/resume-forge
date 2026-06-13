"""文件上传 + AI 提取路由"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from app.database import get_db
from app.models.info_entry import InfoEntry
from app.models.user import User
from app.auth.auth_utils import get_current_user
from app.services.extract_engine import ExtractEngine
import os

router = APIRouter()
extract_engine = ExtractEngine()


@router.post('/upload')
async def upload_and_extract(
    file: UploadFile = File(...),
    category: str = 'work_experience',
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    content = await file.read()
    text = await _extract_text(content, file.filename or '')
    if not text:
        raise HTTPException(400, '无法从文件中提取文本')

    try:
        items = await extract_engine.extract(text, db, current_user.id)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f'AI 提取失败: {str(e)}')

    created = []
    for item in items:
        # Safe category conversion
        cat_raw = item.get('category', category)
        try:
            from app.models.info_entry import InfoCategory
            entry_category = InfoCategory(cat_raw)
        except (ValueError, TypeError):
            entry_category = InfoCategory.WORK_EXPERIENCE

        entry = InfoEntry(
            user_id=current_user.id,
            category=entry_category,
            title=item.get('title', ''),
            company=item.get('company'),
            content=item.get('content', ''),
            raw_input=text[:2000],
            source_file=file.filename,
            start_date=_parse_date(item.get('start_date')),
            end_date=_parse_date(item.get('end_date')),
        )
        db.add(entry)
        created.append(entry)

    await db.commit()
    return {'extracted': len(created), 'items': [{'title': c.title, 'category': c.category.value} for c in created]}



def _parse_date(s):
    if not s:
        return None
    s = str(s).strip()
    if not s:
        return None
    import datetime
    if isinstance(s, (datetime.date, datetime.datetime)):
        return s if isinstance(s, datetime.date) else s.date()
    for fmt in ["%Y-%m-%d", "%Y-%m", "%Y", "%Y.%m", "%Y/%m", "%d/%m/%Y", "%m/%d/%Y"]:
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except (ValueError, TypeError):
            pass
    return None

async def _extract_text(content: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        import pdfplumber, io
        parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: parts.append(t)
        return '\n'.join(parts)
    elif ext in ('.docx', '.doc'):
        import docx, io
        doc = docx.Document(io.BytesIO(content))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return '\n'.join(parts)
    elif ext in ('.xlsx', '.xls'):
        import openpyxl, io
        wb = openpyxl.load_workbook(io.BytesIO(content))
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                rows.append('\t'.join(str(c) if c else '' for c in row))
        return '\n'.join(rows)
    else:
        return content.decode('utf-8', errors='replace')
